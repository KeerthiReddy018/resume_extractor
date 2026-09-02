"""
Resume Information Extraction System
-------------------------------------
100% rule-based / regex / heuristic extraction. NO external LLM or
Generative AI API is used or contacted at any point. Resume content
never leaves the local process.

Supports .pdf and .docx resumes.
Extracts (mandatory): Full Name, Email, Phone, Skills
Extracts (bonus):     Education, Work Experience, LinkedIn, GitHub
"""

import re
import io
import json
from datetime import datetime

import pdfplumber
from pypdf import PdfReader
import docx


# --------------------------------------------------------------------------- #
# 1. TEXT EXTRACTION (PDF / DOCX -> plain text, preserving line breaks)
# --------------------------------------------------------------------------- #

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber, falling back to pypdf."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        text = "\n".join(text_parts)
        if text.strip():
            return text
    except Exception:
        pass

    # Fallback: pypdf (handles some PDFs pdfplumber chokes on)
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(text_parts)
    except Exception:
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX, including paragraphs and table cells."""
    document = docx.Document(io.BytesIO(file_bytes))
    lines = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)

    return "\n".join(lines)


def get_resume_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx",):
        return extract_text_from_docx(file_bytes)
    elif ext == "doc":
        raise ValueError("Legacy .doc format is not supported. Please use .docx or .pdf.")
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Only PDF and DOCX are supported.")


# --------------------------------------------------------------------------- #
# 2. REGEX PATTERNS
# --------------------------------------------------------------------------- #

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")

# Matches national & international phone numbers in various common formats:
# +91 9876543210 | (123) 456-7890 | 123-456-7890 | 123.456.7890 | 9876543210
PHONE_RE = re.compile(
    r"(?:(?:\+|00)\d{1,3}[\s\-.]?)?(?:\(\d{2,4}\)[\s\-.]?)?\d{3,5}[\s\-.]?\d{3,4}[\s\-.]?\d{0,4}"
)

LINKEDIN_RE = re.compile(r"(https?://)?(www\.)?linkedin\.com/[a-zA-Z0-9\-_/%]+", re.IGNORECASE)
GITHUB_RE = re.compile(r"(https?://)?(www\.)?github\.com/[a-zA-Z0-9\-_/%]+", re.IGNORECASE)

SECTION_HEADERS = {
    "education": ["education", "academic background", "academic qualifications", "qualifications"],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment history", "work history", "career history",
    ],
    "skills": ["skills", "technical skills", "core competencies", "key skills", "skill set"],
    "projects": ["projects", "personal projects", "academic projects"],
    "summary": ["summary", "objective", "profile", "about me"],
    "certifications": ["certifications", "certificates", "licenses"],
}

DEGREE_KEYWORDS = [
    "b.tech", "btech", "b.e.", "m.tech", "mtech", "m.e.", "bachelor", "master",
    "b.sc", "bsc", "m.sc", "msc", "phd", "ph.d", "mba", "bca", "mca", "b.com", "bcom",
    "m.com", "mcom", "diploma", "associate degree", "high school", "12th grade", "10th grade",
    "b.a.", "m.a.",
]
# Compiled as whole-word/phrase patterns to avoid false positives like
# "Mumbai" containing the substring "mba".
DEGREE_KEYWORD_RE = re.compile(
    r"(?<![a-zA-Z])(" + "|".join(re.escape(k) for k in DEGREE_KEYWORDS) + r")(?![a-zA-Z])",
    re.IGNORECASE,
)


def _line_has_degree_keyword(line: str) -> bool:
    return bool(DEGREE_KEYWORD_RE.search(line))

# A broad, extensible technical/professional skills dictionary used for matching.
SKILLS_DB = [
    # Programming languages
    "python", "java", "javascript", "typescript", "c++", "c#", "c", "go", "golang",
    "ruby", "php", "swift", "kotlin", "rust", "scala", "r", "matlab", "perl", "dart",
    "objective-c", "shell scripting", "bash",
    # Web
    "html", "css", "html5", "css3", "react", "react.js", "angular", "vue", "vue.js",
    "next.js", "nuxt.js", "node.js", "express.js", "django", "flask", "fastapi",
    "spring", "spring boot", "asp.net", ".net", "jquery", "bootstrap", "tailwind css",
    "graphql", "rest api", "restful apis", "webpack",
    # Data / ML
    "machine learning", "deep learning", "data science", "data analysis",
    "natural language processing", "nlp", "computer vision", "artificial intelligence",
    "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy", "scipy",
    "matplotlib", "seaborn", "opencv", "spacy", "nltk", "power bi", "tableau",
    "data visualization", "statistics", "big data", "hadoop", "spark", "pyspark",
    "etl", "data engineering", "data warehousing", "airflow",
    # Databases
    "sql", "mysql", "postgresql", "postgres", "mongodb", "sqlite", "oracle",
    "microsoft sql server", "redis", "cassandra", "dynamodb", "firebase", "elasticsearch",
    "nosql",
    # Cloud / DevOps
    "aws", "amazon web services", "azure", "gcp", "google cloud platform", "docker",
    "kubernetes", "terraform", "ansible", "jenkins", "ci/cd", "linux", "unix",
    "git", "gitlab", "bitbucket", "devops", "microservices", "nginx",
    "cloudformation", "serverless",
    # Mobile
    "android", "ios", "react native", "flutter", "xamarin",
    # Other tools / soft
    "excel", "microsoft office", "jira", "confluence", "agile", "scrum", "kanban",
    "project management", "communication", "leadership", "teamwork", "problem solving",
    "time management", "critical thinking", "adaptability", "creativity",
    "figma", "adobe photoshop", "adobe illustrator", "ui/ux", "ui design", "ux design",
    "salesforce", "sap", "erp", "seo", "digital marketing", "content writing",
    "machine learning operations", "mlops", "blockchain", "cybersecurity",
    "penetration testing", "networking", "system design", "api development",
    "unit testing", "test automation", "selenium", "postman",
]
# Sort longest-first so multi-word skills match before their substrings
SKILLS_DB_SORTED = sorted(set(SKILLS_DB), key=len, reverse=True)


# --------------------------------------------------------------------------- #
# 3. FIELD EXTRACTORS
# --------------------------------------------------------------------------- #

def extract_email(text: str):
    match = EMAIL_RE.search(text)
    return match.group(0).strip() if match else None


def extract_phone(text: str):
    # Restrict search to lines unlikely to be dates/years to reduce false positives
    candidates = []
    for match in PHONE_RE.finditer(text):
        raw = match.group(0)
        digits = re.sub(r"\D", "", raw)
        # A plausible phone number has 7-15 digits
        if 7 <= len(digits) <= 15:
            candidates.append(raw.strip())
    if not candidates:
        return None
    # Prefer the one with a leading + (country code) if present
    for c in candidates:
        if c.strip().startswith("+"):
            return c.strip()
    return candidates[0]


def extract_linkedin(text: str):
    match = LINKEDIN_RE.search(text)
    if not match:
        return None
    url = match.group(0)
    if not url.lower().startswith("http"):
        url = "https://" + url
    return url.rstrip(").,;")


def extract_github(text: str):
    match = GITHUB_RE.search(text)
    if not match:
        return None
    url = match.group(0)
    if not url.lower().startswith("http"):
        url = "https://" + url
    return url.rstrip(").,;")


def extract_name(text: str, email: str = None):
    """
    Heuristic name extraction:
    1. Look at the first several non-empty lines (resumes almost always
       lead with the candidate's name).
    2. Skip lines that look like headers, emails, phones, urls, or section titles.
    3. Prefer a short line (2-4 words), Title Case or ALL CAPS, with only
       letters/spaces/periods.
    4. Fallback: derive from the email's local-part.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    bad_tokens = ["resume", "curriculum vitae", "cv", "@", "http", "www.",
                  "phone", "email", "address", "objective", "summary"]

    for line in lines[:10]:
        low = line.lower()
        if any(tok in low for tok in bad_tokens):
            continue
        if EMAIL_RE.search(line) or LINKEDIN_RE.search(line) or GITHUB_RE.search(line):
            continue
        if re.search(r"\d", line):
            continue
        words = line.split()
        if 1 <= len(words) <= 4 and all(re.match(r"^[A-Za-z.'\-]+$", w) for w in words):
            # Looks like a plausible name
            return " ".join(w.capitalize() if w.isupper() or w.islower() else w for w in words)

    # Fallback: derive a "guess" from email local-part, e.g. john.doe@x.com -> John Doe
    if email:
        local = email.split("@")[0]
        local = re.sub(r"[._\-]+", " ", local)
        local = re.sub(r"\d+", "", local).strip()
        if local:
            return " ".join(w.capitalize() for w in local.split())

    return None


def find_section_bounds(lines):
    """
    Scans lines and returns {section_key: (start_idx, end_idx)} based on
    SECTION_HEADERS. end_idx is the index of the next detected header (or EOF).
    """
    header_positions = []  # (line_idx, section_key)
    for idx, line in enumerate(lines):
        clean = line.strip().lower().strip(":").strip()
        if not clean or len(clean) > 40:
            continue
        for key, aliases in SECTION_HEADERS.items():
            for alias in aliases:
                if clean == alias or (clean.startswith(alias) and len(clean) <= len(alias) + 3):
                    header_positions.append((idx, key))
                    break

    bounds = {}
    for i, (idx, key) in enumerate(header_positions):
        end = header_positions[i + 1][0] if i + 1 < len(header_positions) else len(lines)
        bounds.setdefault(key, (idx, end))
    return bounds


def extract_skills(text: str, lines, bounds):
    """
    Two-pass approach:
    1. If a 'skills' section is detected, prioritize matching within it.
    2. Always also scan the full document (skills are often scattered),
       matching against SKILLS_DB with word-boundary-safe regex.
    """
    found = set()
    lower_text = text.lower()

    for skill in SKILLS_DB_SORTED:
        pattern = r"(?<![a-zA-Z0-9])" + re.escape(skill) + r"(?![a-zA-Z0-9])"
        if re.search(pattern, lower_text):
            found.add(skill)

    # Also parse comma/pipe/bullet separated tokens from the Skills section
    # to catch skills not present in our dictionary (kept as a bonus list).
    extra = set()
    if "skills" in bounds:
        start, end = bounds["skills"]
        section_text = "\n".join(lines[start + 1:end])
        tokens = re.split(r"[,|/•\u2022\n]+", section_text)
        for tok in tokens:
            tok = tok.strip(" -\t")
            if not tok or len(tok) > 40:
                continue
            low = tok.lower()
            if low in SKILLS_DB_SORTED or low in found:
                continue
            if re.match(r"^[A-Za-z0-9+.#/\- ]{2,40}$", tok):
                extra.add(tok)

    canonical = {s.title() if not s.isupper() else s for s in found}
    # Normalize a few common ones back to preferred casing
    fix_case = {
        "sql": "SQL", "html": "HTML", "css": "CSS", "html5": "HTML5", "css3": "CSS3",
        "aws": "AWS", "gcp": "GCP", "api development": "API Development",
        "rest api": "REST API", "restful apis": "RESTful APIs", "nlp": "NLP",
        "ui/ux": "UI/UX", "php": "PHP", "sap": "SAP", "seo": "SEO",
        "ci/cd": "CI/CD", "mlops": "MLOps", "erp": "ERP", "devops": "DevOps",
        "node.js": "Node.js", "react.js": "React.js", "vue.js": "Vue.js",
        "next.js": "Next.js", "nuxt.js": "Nuxt.js", "express.js": "Express.js",
        "mongodb": "MongoDB", "postgresql": "PostgreSQL", "mysql": "MySQL",
        "pytorch": "PyTorch", "tensorflow": "TensorFlow", "opencv": "OpenCV",
        "asp.net": "ASP.NET", ".net": ".NET", "jquery": "jQuery",
        "scikit-learn": "scikit-learn", "numpy": "NumPy", "matlab": "MATLAB",
        "javascript": "JavaScript", "typescript": "TypeScript", "github": "GitHub",
        "gitlab": "GitLab", "linkedin": "LinkedIn", "dynamodb": "DynamoDB",
        "graphql": "GraphQL", "power bi": "Power BI", "objective-c": "Objective-C",
        "c++": "C++", "c#": "C#",
    }
    final = set()
    for s in found:
        final.add(fix_case.get(s, s.title()))

    result = sorted(final) + sorted(extra)
    # de-duplicate case-insensitively while preserving order
    seen = set()
    deduped = []
    for item in result:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(item)
    return deduped


def extract_education(lines, bounds):
    """Extract education entries as a list of {degree, institution, year}."""
    entries = []
    year_re = re.compile(r"(19|20)\d{2}")
    date_range_re = re.compile(
        r"^\s*(19|20)\d{2}\s*(-|–|to)\s*((19|20)\d{2}|present|current)\s*$", re.IGNORECASE
    )

    if "education" in bounds:
        start, end = bounds["education"]
        block = lines[start + 1:end]
    else:
        # Fallback: scan whole doc for lines containing degree keywords
        block = lines

    consumed = set()
    i = 0
    while i < len(block):
        line = block[i].strip()
        if line and _line_has_degree_keyword(line) and i not in consumed:
            degree = line
            institution = None
            year = None
            y = year_re.search(line)
            if y:
                year = y.group(0)
            # Look at the next 1-2 lines for institution / year if not already found
            for j in range(i + 1, min(i + 3, len(block))):
                nxt = block[j].strip()
                if not nxt or j in consumed:
                    continue
                if not year:
                    y2 = year_re.search(nxt)
                    if y2:
                        year = y2.group(0)
                is_pure_date_range = bool(date_range_re.match(nxt))
                if institution is None and not _line_has_degree_keyword(nxt) and not is_pure_date_range:
                    institution = nxt
                    consumed.add(j)
                elif is_pure_date_range:
                    consumed.add(j)
                    continue
                else:
                    break
            entries.append({
                "degree": degree.strip(" -•\t"),
                "institution": (institution or "").strip(" -•\t") or None,
                "year": year,
            })
        i += 1

    # De-duplicate
    unique = []
    seen = set()
    for e in entries:
        key = (e["degree"].lower(), (e["institution"] or "").lower())
        if key not in seen:
            seen.add(key)
            unique.append(e)
    return unique


def extract_experience(lines, bounds):
    """Extract work experience as a list of {title_or_line, details[]} blocks."""
    if "experience" not in bounds:
        return []

    start, end = bounds["experience"]
    block = [l for l in lines[start + 1:end] if l.strip()]
    if not block:
        return []

    entries = []
    current = None
    year_range_re = re.compile(
        r"(19|20)\d{2}\s*(-|–|to)\s*((19|20)\d{2}|present|current)", re.IGNORECASE
    )

    for line in block:
        stripped = line.strip()
        is_bullet = stripped.startswith(("-", "•", "*", "\u2022"))
        has_year_range = bool(year_range_re.search(stripped))

        if not is_bullet and (has_year_range or len(stripped.split()) <= 12):
            # Treat as a new role/company header line
            if current:
                entries.append(current)
            current = {"role_or_header": stripped, "details": []}
        else:
            if current is None:
                current = {"role_or_header": None, "details": []}
            current["details"].append(stripped.lstrip("-•*\u2022 ").strip())

    if current:
        entries.append(current)

    return entries[:10]  # cap to avoid runaway output on malformed sections


# --------------------------------------------------------------------------- #
# 4. MAIN ENTRY POINT
# --------------------------------------------------------------------------- #

def parse_resume(filename: str, file_bytes: bytes) -> dict:
    text = get_resume_text(filename, file_bytes)
    if not text or not text.strip():
        raise ValueError(
            "No extractable text was found in this file. "
            "It may be a scanned/image-based document, which this rule-based "
            "system does not OCR."
        )

    lines = [l for l in text.split("\n")]
    bounds = find_section_bounds([l.strip() for l in lines])

    email = extract_email(text)
    phone = extract_phone(text)
    name = extract_name(text, email)
    linkedin = extract_linkedin(text)
    github = extract_github(text)
    skills = extract_skills(text, lines, bounds)
    education = extract_education(lines, bounds)
    experience = extract_experience(lines, bounds)

    result = {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": skills,
        "education": education,
        "experience": experience,
        "linkedin": linkedin,
        "github": github,
        "_meta": {
            "source_file": filename,
            "extracted_at": datetime.utcnow().isoformat() + "Z",
            "characters_extracted": len(text),
        },
    }
    return result


if __name__ == "__main__":
    import sys
    with open(sys.argv[1], "rb") as f:
        data = f.read()
    print(json.dumps(parse_resume(sys.argv[1].split("/")[-1], data), indent=2))
